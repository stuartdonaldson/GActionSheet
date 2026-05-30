"""
Unit tests for scn/surfaces.py — DocReader, SheetReader, TrackerReader (GTaskSheet-5vwu.6).

Pure unit — no network, no GAS, no local.settings.json required.
Covers every AC from the bead:
  - Sheet reads are docId-scoped (no cross-session pollution)
  - Tracker rows expose chip assignee email
  - Doc reader exposes all occurrences for identical-occurrence checks
"""
import io
import re
import pytest
from lxml import etree

import docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import openpyxl

from scn.ai import ai
from scn.surfaces import DocReader, SheetReader, TrackerReader


# ---------------------------------------------------------------------------
# Docx fixture helpers
# ---------------------------------------------------------------------------

_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_MAILTO_PREFIX = "mailto:"


def _add_hyperlink_run(para, display_text: str, url: str) -> None:
    """Add a hyperlink run to an existing paragraph, registering the relationship."""
    part = para.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(f"{{{_R_NS}}}id", r_id)

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = display_text
    t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    run.append(t)
    hyperlink.append(run)
    para._element.append(hyperlink)


def _add_plain_run(para, text: str) -> None:
    run = para.add_run(text)


def _make_docx_bytes(paragraphs: list[str | tuple]) -> bytes:
    """
    Build an in-memory docx from a list of paragraph specs.

    Each spec is either:
      - str: plain text paragraph
      - tuple of (prefix_text, chip_email, chip_display, suffix_text)
        builds: <run prefix_text><hyperlink mailto:chip_email chip_display><run suffix_text>
    """
    doc = Document()
    for spec in paragraphs:
        para = doc.add_paragraph()
        # Remove the auto-added run from add_paragraph (it's empty by default)
        for run in para.runs:
            run._element.getparent().remove(run._element)

        if isinstance(spec, str):
            _add_plain_run(para, spec)
        else:
            prefix, email, display, suffix = spec
            if prefix:
                _add_plain_run(para, prefix)
            _add_hyperlink_run(para, display, f"{_MAILTO_PREFIX}{email}")
            if suffix:
                _add_plain_run(para, suffix)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_tracker_docx_bytes(rows: list[dict], section_heading: str = "=== Tracked Actions ===") -> bytes:
    """
    Build a docx with an '=== Tracked Actions ===' section and a tracker table.

    Each row dict maps header name → cell spec (str or (email, display_name) tuple for chip).
    HEADERS = ['ID', 'Assignee Email', 'Assignee Name', 'Action', 'Status', 'Date Created', 'Date Modified']
    Pass a (email, display_name) tuple for 'Assignee Name' to create a chip hyperlink.
    """
    HEADERS = ['ID', 'Assignee Email', 'Assignee Name', 'Action', 'Status', 'Date Created', 'Date Modified']

    doc = Document()
    doc.add_paragraph(section_heading)

    table = doc.add_table(rows=1 + len(rows), cols=len(HEADERS))
    # header row
    for c, h in enumerate(HEADERS):
        table.rows[0].cells[c].text = h
    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[1 + r_idx]
        for c, h in enumerate(HEADERS):
            spec = row_data.get(h, "")
            cell = row.cells[c]
            if isinstance(spec, tuple):
                # chip: (email, display_name)
                email, display = spec
                # Clear default paragraph then add hyperlink
                cell.paragraphs[0]._element.clear()
                para = cell.paragraphs[0]
                _add_hyperlink_run(para, display, f"{_MAILTO_PREFIX}{email}")
            else:
                cell.text = str(spec) if spec else ""

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# DocReader tests
# ---------------------------------------------------------------------------

class TestDocReader:
    def test_email_parsed_path(self):
        """AI-N: followed by email in text → assignee_source='parsed'."""
        docx_bytes = _make_docx_bytes([
            "AI-1: user@example.com Do the thing (Open)"
        ])
        results = DocReader().read(docx_bytes, "any-doc-id")
        assert len(results) == 1
        r = results[0]
        assert r.action_id == "AI-1"
        assert r.assignee == "user@example.com"
        assert r.action == "Do the thing"
        assert r.status == "Open"
        assert r.assignee_source == "parsed"

    def test_chip_path(self):
        """AI-N: followed by person chip → assignee_source='chip'."""
        docx_bytes = _make_docx_bytes([
            ("AI-2: ", "chip@example.com", "Chip User", " Do the chip thing (In Progress)")
        ])
        results = DocReader().read(docx_bytes, "any-doc-id")
        assert len(results) == 1
        r = results[0]
        assert r.action_id == "AI-2"
        assert r.assignee == "chip@example.com"
        assert r.action == "Do the chip thing"
        assert r.status == "In Progress"
        assert r.assignee_source == "chip"

    def test_no_assignee(self):
        """AI-N: with no assignee → assignee=None, assignee_source=None."""
        docx_bytes = _make_docx_bytes([
            "AI-3: Do the thing"
        ])
        results = DocReader().read(docx_bytes, "any-doc-id")
        assert len(results) == 1
        r = results[0]
        assert r.action_id == "AI-3"
        assert r.assignee is None
        assert r.assignee_source is None
        assert r.action == "Do the thing"
        assert r.status is None

    def test_non_matching_paragraph_skipped(self):
        """Paragraph without AI-N: token not returned."""
        docx_bytes = _make_docx_bytes([
            "Just a regular paragraph",
            "Another one without token",
            "AI-1: user@example.com An action (Open)",
        ])
        results = DocReader().read(docx_bytes, "any-doc-id")
        assert len(results) == 1
        assert results[0].action_id == "AI-1"

    def test_multiple_occurrences_same_id_both_returned(self):
        """Two paragraphs with the same action_id are both returned (engine checks identity)."""
        docx_bytes = _make_docx_bytes([
            "AI-1: user@example.com Do the thing (Open)",
            "AI-1: user@example.com Do the thing (Open)",
        ])
        results = DocReader().read(docx_bytes, "any-doc-id")
        assert len(results) == 2
        assert all(r.action_id == "AI-1" for r in results)

    def test_status_absent_returns_none(self):
        """No trailing (Status) token → status=None."""
        docx_bytes = _make_docx_bytes([
            "AI-5: Do the thing"
        ])
        results = DocReader().read(docx_bytes, "any-doc-id")
        assert results[0].status is None

    def test_status_present_extracted(self):
        """Trailing (In Progress) extracted correctly."""
        docx_bytes = _make_docx_bytes([
            "AI-6: Do the thing (In Progress)"
        ])
        results = DocReader().read(docx_bytes, "any-doc-id")
        assert results[0].status == "In Progress"

    def test_multiple_actions_in_one_doc(self):
        """Multiple different actions are all returned."""
        docx_bytes = _make_docx_bytes([
            "AI-1: a@x.com First action (Open)",
            "Not an action",
            "AI-2: b@x.com Second action (Done)",
        ])
        results = DocReader().read(docx_bytes, "any-doc-id")
        assert len(results) == 2
        assert {r.action_id for r in results} == {"AI-1", "AI-2"}

    def test_doc_id_param_unused_for_doc_surface(self):
        """DocReader returns all AI-N paragraphs regardless of doc_id parameter (no filtering)."""
        docx_bytes = _make_docx_bytes(["AI-1: Do the thing"])
        assert len(DocReader().read(docx_bytes, "doc-abc")) == 1
        assert len(DocReader().read(docx_bytes, "doc-xyz")) == 1


# ---------------------------------------------------------------------------
# SheetReader tests
# ---------------------------------------------------------------------------

_TARGET_DOC_ID = "1AbCdEfGhIjKlMnOpQrStUvWxYz"
_OTHER_DOC_ID = "9ZzZzZzZzZzZzZzZzZzZzZzZzZz"
_TARGET_DOC_URL = f"https://docs.google.com/document/d/{_TARGET_DOC_ID}/edit"
_OTHER_DOC_URL = f"https://docs.google.com/document/d/{_OTHER_DOC_ID}/edit"


def _make_xlsx_bytes(rows: list[dict]) -> bytes:
    """
    Build xlsx with SHEET_HEADERS header row and given data rows.

    Row dict keys: global_id, action_id, assignee_email, assignee_name,
                   action_text, status, document_url, doc_name,
                   created_date, modified_date, sync_status
    document_url → stored as =HYPERLINK("url","name") formula in col 7.
    """
    HEADERS = [
        "NamedRangeId", "ID", "Assignee Email", "Assignee Name",
        "Action", "Status", "Document",
        "Date Created", "Date Modified", "Sync Status"
    ]
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(HEADERS)
    for row in rows:
        doc_url = row.get("document_url", "")
        doc_name = row.get("doc_name", "Doc")
        formula = f'=HYPERLINK("{doc_url}","{doc_name}")' if doc_url else ""
        ws.append([
            row.get("global_id", ""),
            row.get("action_id", ""),
            row.get("assignee_email", ""),
            row.get("assignee_name", ""),
            row.get("action_text", ""),
            row.get("status", ""),
            formula,
            row.get("created_date", ""),
            row.get("modified_date", ""),
            row.get("sync_status", ""),
        ])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


class TestSheetReader:
    def test_row_matching_doc_id_returned(self):
        """Row whose document_formula contains the target doc_id is returned."""
        xlsx_bytes = _make_xlsx_bytes([{
            "global_id": f"{_TARGET_DOC_ID}/AI-1",
            "action_id": "AI-1",
            "assignee_email": "user@x.com",
            "assignee_name": "User Name",
            "action_text": "Do the thing",
            "status": "Open",
            "document_url": _TARGET_DOC_URL,
            "doc_name": "My Doc",
            "sync_status": "Synced",
        }])
        results = SheetReader().read(xlsx_bytes, _TARGET_DOC_ID)
        assert len(results) == 1
        r = results[0]
        assert r.action == "Do the thing"
        assert r.action_id == "AI-1"
        assert r.assignee == "user@x.com"
        assert r.status == "Open"

    def test_row_different_doc_id_filtered_out(self):
        """Row for a different doc is excluded — no cross-session pollution (AC)."""
        xlsx_bytes = _make_xlsx_bytes([{
            "global_id": f"{_OTHER_DOC_ID}/AI-1",
            "action_id": "AI-1",
            "assignee_email": "other@x.com",
            "action_text": "Other doc action",
            "document_url": _OTHER_DOC_URL,
            "doc_name": "Other Doc",
        }])
        results = SheetReader().read(xlsx_bytes, _TARGET_DOC_ID)
        assert results == []

    def test_mixed_rows_only_target_returned(self):
        """Two rows with different doc_ids — only the target doc row is returned."""
        xlsx_bytes = _make_xlsx_bytes([
            {
                "action_id": "AI-1", "action_text": "Target action",
                "document_url": _TARGET_DOC_URL, "doc_name": "Target",
            },
            {
                "action_id": "AI-2", "action_text": "Other action",
                "document_url": _OTHER_DOC_URL, "doc_name": "Other",
            },
        ])
        results = SheetReader().read(xlsx_bytes, _TARGET_DOC_ID)
        assert len(results) == 1
        assert results[0].action == "Target action"

    def test_global_id_attribute_present(self):
        """Returned ai carries global_id as a dynamic attribute."""
        global_id = f"{_TARGET_DOC_ID}/AI-1"
        xlsx_bytes = _make_xlsx_bytes([{
            "global_id": global_id, "action_id": "AI-1",
            "action_text": "Do it", "document_url": _TARGET_DOC_URL,
        }])
        results = SheetReader().read(xlsx_bytes, _TARGET_DOC_ID)
        assert hasattr(results[0], "global_id")
        assert results[0].global_id == global_id

    def test_assignee_name_attribute_present(self):
        """Returned ai carries assignee_name as a dynamic attribute."""
        xlsx_bytes = _make_xlsx_bytes([{
            "action_id": "AI-1", "assignee_name": "Test User",
            "action_text": "Do it", "document_url": _TARGET_DOC_URL,
        }])
        results = SheetReader().read(xlsx_bytes, _TARGET_DOC_ID)
        assert hasattr(results[0], "assignee_name")
        assert results[0].assignee_name == "Test User"

    def test_sync_status_attribute_present(self):
        """Returned ai carries sync_status as a dynamic attribute."""
        xlsx_bytes = _make_xlsx_bytes([{
            "action_id": "AI-1", "action_text": "Do it",
            "sync_status": "Dirty", "document_url": _TARGET_DOC_URL,
        }])
        results = SheetReader().read(xlsx_bytes, _TARGET_DOC_ID)
        assert hasattr(results[0], "sync_status")
        assert results[0].sync_status == "Dirty"

    def test_doc_id_derived_attribute(self):
        """Returned ai carries doc_id attribute derived from formula URL."""
        xlsx_bytes = _make_xlsx_bytes([{
            "action_id": "AI-1", "action_text": "Do it",
            "document_url": _TARGET_DOC_URL, "doc_name": "My Doc",
        }])
        results = SheetReader().read(xlsx_bytes, _TARGET_DOC_ID)
        assert hasattr(results[0], "doc_id")
        assert results[0].doc_id == _TARGET_DOC_ID

    def test_doc_name_derived_attribute(self):
        """Returned ai carries doc_name attribute derived from formula."""
        xlsx_bytes = _make_xlsx_bytes([{
            "action_id": "AI-1", "action_text": "Do it",
            "document_url": _TARGET_DOC_URL, "doc_name": "My Doc",
        }])
        results = SheetReader().read(xlsx_bytes, _TARGET_DOC_ID)
        assert hasattr(results[0], "doc_name")
        assert results[0].doc_name == "My Doc"

    def test_assignee_source_not_set_for_sheet(self):
        """Sheet rows carry no assignee_source (chip concept not applicable)."""
        xlsx_bytes = _make_xlsx_bytes([{
            "action_id": "AI-1", "action_text": "Do it",
            "document_url": _TARGET_DOC_URL,
        }])
        results = SheetReader().read(xlsx_bytes, _TARGET_DOC_ID)
        assert results[0].assignee_source is None

    def test_empty_sheet_returns_empty_list(self):
        """Sheet with no data rows returns empty list."""
        xlsx_bytes = _make_xlsx_bytes([])
        results = SheetReader().read(xlsx_bytes, _TARGET_DOC_ID)
        assert results == []


# ---------------------------------------------------------------------------
# TrackerReader tests
# ---------------------------------------------------------------------------

class TestTrackerReader:
    def test_chip_row_assignee_source_chip(self):
        """Tracker row with chip in Assignee Name cell → assignee_source='chip'."""
        docx_bytes = _make_tracker_docx_bytes([{
            "ID": "AI-1",
            "Assignee Email": "",
            "Assignee Name": ("chip@example.com", "Chip User"),  # chip tuple
            "Action": "Do the tracker thing",
            "Status": "Open",
        }])
        results = TrackerReader().read(docx_bytes, "any-doc-id")
        assert len(results) == 1
        r = results[0]
        assert r.action_id == "AI-1"
        assert r.assignee == "chip@example.com"
        assert r.assignee_source == "chip"
        assert r.action == "Do the tracker thing"
        assert r.status == "Open"

    def test_plain_text_row_assignee_source_parsed(self):
        """Tracker row with plain text email → assignee_source='parsed'."""
        docx_bytes = _make_tracker_docx_bytes([{
            "ID": "AI-2",
            "Assignee Email": "plain@example.com",
            "Assignee Name": "Plain User",
            "Action": "Another action",
            "Status": "In Progress",
        }])
        results = TrackerReader().read(docx_bytes, "any-doc-id")
        assert len(results) == 1
        r = results[0]
        assert r.assignee == "plain@example.com"
        assert r.assignee_source == "parsed"
        assert r.action == "Another action"

    def test_chip_row_exposes_assignee_email(self):
        """AC: tracker rows expose chip assignee email."""
        docx_bytes = _make_tracker_docx_bytes([{
            "ID": "AI-3",
            "Assignee Name": ("minister@northlakeuu.org", "Northlake Minister"),
            "Action": "Chair the meeting",
            "Status": "Open",
        }])
        results = TrackerReader().read(docx_bytes, "any-doc-id")
        assert results[0].assignee == "minister@northlakeuu.org"

    def test_assignee_name_attribute_present(self):
        """Returned ai carries assignee_name dynamic attribute."""
        docx_bytes = _make_tracker_docx_bytes([{
            "ID": "AI-1",
            "Assignee Name": ("chip@x.com", "Chip Name"),
            "Action": "Do it",
            "Status": "Open",
        }])
        results = TrackerReader().read(docx_bytes, "any-doc-id")
        assert hasattr(results[0], "assignee_name")
        assert results[0].assignee_name == "Chip Name"

    def test_no_section_heading_returns_empty(self):
        """No tracker section → empty list."""
        doc = Document()
        doc.add_paragraph("Just a regular document")
        buf = io.BytesIO()
        doc.save(buf)
        results = TrackerReader().read(buf.getvalue(), "any-doc-id")
        assert results == []

    def test_multiple_tracker_rows(self):
        """Multiple data rows all returned."""
        docx_bytes = _make_tracker_docx_bytes([
            {"ID": "AI-1", "Assignee Name": ("a@x.com", "A"), "Action": "Act 1", "Status": "Open"},
            {"ID": "AI-2", "Assignee Name": ("b@x.com", "B"), "Action": "Act 2", "Status": "Done"},
        ])
        results = TrackerReader().read(docx_bytes, "any-doc-id")
        assert len(results) == 2
        assert {r.action_id for r in results} == {"AI-1", "AI-2"}

    def test_doc_id_param_unused_for_tracker(self):
        """TrackerReader returns all rows regardless of doc_id param (tracker is per-doc)."""
        docx_bytes = _make_tracker_docx_bytes([{
            "ID": "AI-1", "Assignee Email": "u@x.com", "Action": "Do it", "Status": "Open"
        }])
        assert len(TrackerReader().read(docx_bytes, "doc-abc")) == 1
        assert len(TrackerReader().read(docx_bytes, "doc-xyz")) == 1
