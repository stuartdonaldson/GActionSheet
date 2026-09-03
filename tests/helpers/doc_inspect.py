"""python-docx helpers for asserting tracked-actions table and floating-action content.

``floating_actions`` is the **independent track** of the twin-track oracle
(gts-e5cl): a Python parse of ADR-0027's floating-action paragraph grammar,
read from the written spec —
``knowledge-base/adr/0027-floating-action-paragraph-grammar.md`` and
``docs/interfaces/action-portable-text.md`` — and deliberately not from
``src/SyncManager.js``, ``src/PortableText.js``, ``scripts/apt_lib.py`` or any
blessed ``tests/fixtures/*.apt.txt`` corpus.  A parser taught by the system it
judges cannot contradict that system, which is exactly how 20 of 21 unscanned
actions were frozen as expected output on 2026-08-29.

What it does NOT cover, by design: inline bold/italic/link *runs* inside an
action body or a field value.  Field values are returned as plain joined
text, not ADR-0027 rule 15's ``{text, runs}`` -- gts-dxgo (stage
`apt-presentation`) added ``has_status_icon`` (a flush-inserted inline
image's presence) and left rule 15 run extraction for a follow-up bead; see
its stage handoff.
"""
import io
import re
from dataclasses import dataclass, field as _dc_field

import docx
from docx.oxml.ns import qn

_SECTION_HEADING = "=== Tracked Actions ==="
_EMAIL_RE = re.compile(r'[\w.+\-]+@[\w\-]+(?:\.[a-z]{2,})+', re.IGNORECASE)

_R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'

# ADR-0027 rule 1 / ADR-0023: 'ACT-N:' canonical on write, 'AI-N:' read-valid.
_TOKEN_RE = re.compile(r'(ACT|AI)-(\d+):[ \t]*')
# gts-lu5k: same token, but counted per PHYSICAL LINE within a paragraph (line
# start or right after a soft-return '\n') — mirrors src/SyncManager.js's
# _collectActionsFromParagraph dispatch (tokenCount via '(?:^|\n)TOKEN').
_TOKEN_COUNT_RE = re.compile(r'(?:^|\n)(?:ACT|AI)-\d+:')
_TOKEN_PREFIX_RE = re.compile(r'(ACT|AI)-\d+')       # rule 6: looks like an action…
_BARE_TRIGGER_RE = re.compile(r'(ACT|AI):[ \t]*')    # unnumbered trigger, pre-flush
_ASSIGNEE_RE = re.compile(r'@?(' + _EMAIL_RE.pattern + r')[ \t]*', re.IGNORECASE)
_STATUS_GROUP_RE = re.compile(r'\(([^)]*)\)')

# ADR-0027 rule 5: every space-separated word Title Case, ≤32 chars total,
# then a colon followed by a space, a tab, or the end of the line.
_FIELD_LINE_RE = re.compile(
    r'^([A-Z][A-Za-z0-9_-]*(?: [A-Z][A-Za-z0-9_-]*)*):(?:([ \t])(.*))?$'
)
_FIELD_NAME_MAX = 32

_CHIP = '\x00'  # stands in for a PERSON chip while the header line is parsed

UNPARSEABLE = 'unparseable-action-paragraph'


@dataclass
class ParsedAction:
    """One floating-action paragraph, as the grammar reads it."""

    token: str | None = None            # 'ACT-9' / 'AI-9'; None for a bare trigger
    prefix: str | None = None           # 'ACT' | 'AI'
    number: int | None = None
    assignee_email: str | None = None
    assignee_name: str = ''             # chip display name; '' for a text email
    assignee_source: str | None = None  # 'chip' | 'text' | None
    action_text: str = ''               # header body + prose lines, joined with '\n'
    status: str | None = None           # None when the doc carries no token
    has_explicit_status: bool = False
    custom_fields: dict[str, str] = _dc_field(default_factory=dict)
    token_linked: bool = False          # does the ACT-N: header carry its chip link?
    token_url: str | None = None
    pending: bool = False               # bare 'AI:'/'ACT:' trigger, not yet assigned
    container: str = 'body'             # 'body' | 'list-item' | 'table-cell'
    body_index: int = -1                # paragraph index in document order
    raw_text: str = ''
    error: str | None = None            # UNPARSEABLE, or None
    has_status_icon: bool = False       # a flush-inserted inline image is present


def load_doc(docx_bytes: bytes) -> docx.Document:
    return docx.Document(io.BytesIO(docx_bytes))


# ---------------------------------------------------------------------------
# Low-level text extraction
# ---------------------------------------------------------------------------

def _local(tag: str) -> str:
    return tag.split('}')[-1] if '}' in tag else tag


def _node_text(node) -> str:
    """Text under ``node``, keeping a soft return as '\\n' and a tab as '\\t'.

    ``para.text`` concatenates only ``w:t`` runs, so it silently drops the
    ``<w:br/>`` a Shift+Enter exports as and the ``<w:tab/>`` a literal tab
    exports as — which makes ADR-0027's ``SR`` terminal and rule 8's field-label
    tab invisible, the two separators the whole continuation grammar keys off.
    """
    parts = []
    for el in node.iter():
        tag = _local(el.tag)
        if tag == 't' and el.text:
            parts.append(el.text)
        elif tag == 'br' and el.get(qn('w:type')) != 'page':
            parts.append('\n')
        elif tag == 'tab':
            parts.append('\t')
    return ''.join(parts).replace('\v', '\n')


def _rel_maps(document_part) -> tuple[dict[str, str], dict[str, str]]:
    """Return (rId -> email, rId -> url) for a part's hyperlink relationships."""
    import urllib.parse as _urlparse
    emails: dict[str, str] = {}
    urls: dict[str, str] = {}
    try:
        rels = document_part.rels.values()
    except Exception:
        return emails, urls
    for rel in rels:
        if 'hyperlink' not in rel.reltype:
            continue
        url = rel.target_ref or ''
        urls[rel.rId] = url
        m = re.match(r'mailto:([^\?&]+)', url, re.IGNORECASE)
        if m:
            emails[rel.rId] = m.group(1).strip()
            continue
        m = re.search(r'[?&]email=([^&]+)', url, re.IGNORECASE)
        if m:
            emails[rel.rId] = _urlparse.unquote(m.group(1)).strip()
    return emails, urls


def _segments(p_el, emails: dict[str, str], urls: dict[str, str]) -> tuple[str, list, list]:
    """Flatten a ``w:p`` to (text, span table, chips).

    A PERSON chip contributes a single ``_CHIP`` sentinel character instead of
    its display text, so the assignee's *position* stays visible to the header
    line parse while its display name stays out of the action text.  The span
    table is a list of ``(start, end, url)`` for hyperlink URL lookup by offset.
    """
    out: list[str] = []
    spans: list[tuple[int, int, str | None]] = []
    chips: list[tuple[str, str]] = []
    pos = 0
    for child in p_el:
        tag = _local(child.tag)
        if tag == 'hyperlink':
            r_id = child.get(f'{{{_R_NS}}}id', '')
            email = emails.get(r_id)
            if email:
                chips.append((email, _node_text(child)))
                piece, url = _CHIP, None
            else:
                piece, url = _node_text(child), urls.get(r_id)
        elif tag in ('pPr', 'bookmarkStart', 'bookmarkEnd', 'proofErr'):
            continue
        else:
            piece, url = _node_text(child), None
        if not piece:
            continue
        out.append(piece)
        spans.append((pos, pos + len(piece), url))
        pos += len(piece)
    return ''.join(out), spans, chips


def _url_at(spans, index: int) -> str | None:
    for start, end, url in spans:
        if start <= index < end:
            return url
    return None


def _has_image(p_el) -> bool:
    """docs/interfaces/action-portable-text.md "Status icon": a flush inserts
    a small inline image at the start of the paragraph, exported as a
    ``w:drawing`` (or legacy ``w:pict``) element inside a run. Presence-only
    -- position isn't checked, since a flush always places it before any
    other content (gts-dxgo's scope is presence, not layout)."""
    for el in p_el.iter():
        if _local(el.tag) in ('drawing', 'pict'):
            return True
    return False


def status_icon_sizes_pt(document: docx.Document) -> list[tuple[float, float]]:
    """gts-bxrt: (width_pt, height_pt) for every inline image in the document,
    document order -- the size oracle for the new 'Status Icon Size' Config
    row, distinct from `_has_image`'s presence-only check above. Uses
    python-docx's own `inline_shapes` (EMU-backed `Length`, `.pt` converts),
    not manual XML walking, since exact dimensions (not just drawing
    presence) are what's under test here."""
    return [
        (shape.width.pt, shape.height.pt)
        for shape in document.inline_shapes
    ]


def _container_of(p_el) -> str:
    node = p_el.getparent()
    while node is not None:
        if _local(node.tag) == 'tc':
            return 'table-cell'
        node = node.getparent()
    if p_el.find('.//' + qn('w:numPr')) is not None:
        return 'list-item'
    return 'body'


# ---------------------------------------------------------------------------
# Grammar
# ---------------------------------------------------------------------------

def _split_status(header: str) -> tuple[str, str | None]:
    """ADR-0027 rule 4 + gts-1tbe: the LAST '(...)' group on the header line.

    It qualifies only when what follows it — ignoring trailing whitespace — is
    empty or begins with a non-word character.  '(draft) proposal' is therefore
    literal text, while '(Open) - done' is a status with its trailing text kept
    (gts-v0py: author-typed text after the token is preserved, never dropped).
    """
    groups = list(_STATUS_GROUP_RE.finditer(header))
    if not groups:
        return header.strip(), None
    m = groups[-1]
    tail = header[m.end():].strip()
    if tail and (tail[0].isalnum() or tail[0] == '_'):
        return header.strip(), None
    before = header[:m.start()].rstrip()
    body = f'{before} {tail}'.strip() if tail else before.strip()
    return body, m.group(1).strip()


def _blocks(header_body: str, continuation: list[str]) -> tuple[str, dict[str, str]]:
    """ADR-0027 rule 5a — prose attaches to whichever block is open."""
    action_lines = [header_body]
    fields: dict[str, list[str]] = {}
    open_field: str | None = None

    for raw in continuation:
        line = raw.lstrip(' \t')  # rule 5: the indent is presentational, never stored
        m = _FIELD_LINE_RE.match(line)
        if m and len(m.group(1)) <= _FIELD_NAME_MAX:
            name = m.group(1)
            open_field = name
            fields.setdefault(name, []).append(m.group(3) if m.group(2) else '')
        elif open_field is not None:
            fields[open_field].append(line)
        else:
            action_lines.append(line)

    return '\n'.join(action_lines), {k: '\n'.join(v) for k, v in fields.items()}


def _finish_record(base: ParsedAction, rest: str, chips, chip_i: int) -> ParsedAction:
    """Shared tail of a token record's parse: assignee, status, action text,
    custom fields. `rest` is everything after the token match (and its
    trailing [ \\t]*) through the end of this record's own text; `chip_i` is
    how many chips (in document order) have already been consumed before
    `rest` begins.
    """
    # assignee — a chip immediately after the token, else an optionally @-sigilled email
    if rest.startswith(_CHIP) and chip_i < len(chips):
        base.assignee_email, base.assignee_name = chips[chip_i]
        base.assignee_source = 'chip'
        chip_i += 1
        rest = rest[1:].lstrip(' \t')
    else:
        am = _ASSIGNEE_RE.match(rest)
        if am:
            base.assignee_email = am.group(1)
            base.assignee_source = 'text'
            rest = rest[am.end():]

    # any further chip is ordinary content — restore its display text
    while _CHIP in rest:
        display = chips[chip_i][1] if chip_i < len(chips) else ''
        rest = rest.replace(_CHIP, display, 1)
        chip_i += 1

    header, _, tail = rest.partition('\n')
    header_body, status = _split_status(header)
    base.status = status
    base.has_explicit_status = status is not None

    continuation = tail.split('\n') if tail else []
    base.action_text, base.custom_fields = _blocks(header_body, continuation)
    return base


def _parse_paragraph(p_el, index: int, emails, urls) -> list[ParsedAction]:
    text, spans, chips = _segments(p_el, emails, urls)
    if not text.strip():
        return []

    container = _container_of(p_el)
    has_status_icon = _has_image(p_el)
    token_count = len(_TOKEN_COUNT_RE.findall(text))

    lead = len(text) - len(text.lstrip(' \t'))
    rest0 = text[lead:]
    single_match = _TOKEN_RE.match(rest0) if token_count == 1 else None

    if token_count == 0 or (token_count == 1 and single_match is None):
        # gts-lu5k: token_count==1-but-not-anchored (a token on a
        # non-first physical line, with no OTHER token in the paragraph)
        # still needs the soft-return walk below, NOT this single-line
        # branch — a lone token buried after an intro line is exactly the
        # shape this bead exists to stop silently dropping. Only a true
        # zero-token paragraph (bare trigger / rule-6 unparseable / plain
        # prose) belongs here.
        if token_count == 1:
            pass  # fall through to the soft-return walk below
        else:
            raw_text = text.replace(_CHIP, '')
            base = ParsedAction(
                container=container, body_index=index,
                raw_text=raw_text, has_status_icon=has_status_icon,
            )
            m = _BARE_TRIGGER_RE.match(rest0)
            if m:
                base.pending = True
                rest = rest0[m.end():]
                return [_finish_record(base, rest, chips, chip_i=0)]
            if _TOKEN_PREFIX_RE.match(rest0):
                # rule 6: begins like an action but does not complete the grammar.
                base.error = UNPARSEABLE
                return [base]
            return []

    if token_count == 1 and single_match is not None:
        # Single token, anchored at the paragraph's own start: existing
        # fast-path record, unchanged.
        m = single_match
        raw_text = text.replace(_CHIP, '')
        base = ParsedAction(
            container=container, body_index=index,
            raw_text=raw_text, has_status_icon=has_status_icon,
        )
        base.prefix, base.number = m.group(1), int(m.group(2))
        base.token = f'{m.group(1)}-{m.group(2)}'
        base.token_url = _url_at(spans, lead)
        base.token_linked = base.token_url is not None
        rest = rest0[m.end():]
        return [_finish_record(base, rest, chips, chip_i=0)]

    # gts-lu5k: multi-token OR a single token that follows an intro line —
    # mirrors src/SyncManager.js's _parseSoftReturnParagraphActions. Walk
    # every '(?:^|\n)TOKEN' occurrence and slice out one record per token,
    # from just after its own match through the character before the next
    # token's match (its own '\n' boundary, if any, excluded from either
    # side) or end of paragraph. Leading context text before the first
    # token is skipped, matching the production scanner (AC-3).
    matches = list(_TOKEN_COUNT_RE.finditer(text))
    results: list[ParsedAction] = []
    for mi, tm in enumerate(matches):
        # tm.group(0) is '(\n)?ACT-N:' — strip the leading '\n' (if any) to
        # find where the token text itself actually starts.
        token_start = tm.start() + (1 if text[tm.start()] == '\n' else 0)
        # gts-lu5k: re-match with _TOKEN_RE (not _TOKEN_COUNT_RE) anchored at
        # token_start so the trailing '[ \t]*' after the colon is consumed
        # the same way the single-token fast path consumes it — otherwise
        # `rest` below starts with the whitespace _ASSIGNEE_RE has no
        # leading-whitespace tolerance for, and a text-email assignee is
        # silently lost.
        tok_m = _TOKEN_RE.match(text, token_start)
        # tok_m always matches: tm was built from the exact same alternation.
        prefix, number = tok_m.group(1), int(tok_m.group(2))
        consumed_end = tok_m.end()

        record_end = matches[mi + 1].start() if mi + 1 < len(matches) else len(text)
        rest = text[consumed_end:record_end]
        raw_text = text[token_start:record_end].replace(_CHIP, '')
        chip_i_before = text[:consumed_end].count(_CHIP)

        base = ParsedAction(
            container=container, body_index=index,
            raw_text=raw_text, has_status_icon=has_status_icon,
        )
        base.prefix, base.number = prefix, number
        base.token = f'{prefix}-{number}'
        base.token_url = _url_at(spans, token_start)
        base.token_linked = base.token_url is not None
        results.append(_finish_record(base, rest, chips, chip_i=chip_i_before))
    return results


def floating_actions(document: docx.Document) -> list[ParsedAction]:
    """Every floating-action paragraph in the document, in document order.

    Detection is the ``ACT-N:``/``AI-N:`` token (ADR-0023's dual prefix), never
    ``w:numPr`` — the grammar has no list-item requirement, so a body paragraph,
    a bullet and a table cell are all reached identically. A paragraph can
    yield more than one record — one per ``ACT-N:``/``AI-N:`` token line
    (gts-lu5k), matching what ``src/SyncManager.js``'s scanner reports for the
    same paragraph.

    Records that are present but not fully established are returned, not
    dropped: an unlinked ``ACT-N:`` header (``token_linked=False``), a bare
    ``AI:`` trigger (``pending=True``), and a paragraph that begins like an
    action but fails the grammar (``error=UNPARSEABLE``, rule 6).
    """
    emails, urls = _rel_maps(document.part)
    results: list[ParsedAction] = []
    for index, p_el in enumerate(document.element.body.iter(qn('w:p'))):
        results.extend(_parse_paragraph(p_el, index, emails, urls))
    return results


def unparseable_paragraphs(document: docx.Document) -> list[ParsedAction]:
    """Rule 6 records only — paragraphs that look like actions but do not parse."""
    return [a for a in floating_actions(document) if a.error]


# ---------------------------------------------------------------------------
# Other document views (unchanged)
# ---------------------------------------------------------------------------

def paragraph_texts_with_breaks(document: docx.Document) -> list[str]:
    """Return one string per body paragraph, with soft line breaks kept as '\\n'.

    Soft returns (``<w:br/>``) and literal tabs (``<w:tab/>``) are preserved as
    '\\n' and '\\t' — see ``_node_text``, which this shares with the grammar
    parser so the two can never disagree about what a paragraph's text is.
    Page breaks are ignored.
    """
    return [_node_text(para._element) for para in document.paragraphs]


def paragraph_bold_text(document: docx.Document) -> list[str]:
    """Return one string per body paragraph, containing ONLY the text of runs
    whose direct w:rPr marks them bold — everything else is dropped.

    gts-po8t: used to assert _renderCustomFieldLines' field-name-is-bolded
    contract (SyncManager.js) — a docx run's bold state lives in
    w:rPr/w:b (or an inherited/toggled style, not resolved here since the
    add-on always sets bold explicitly on the field-line updateTextStyle
    request rather than relying on inheritance).
    """
    texts = []
    for para in document.paragraphs:
        parts = []
        for run in para.runs:
            if run.bold:
                parts.append(run.text or '')
        texts.append(''.join(parts))
    return texts


def tracked_actions_table(document: docx.Document) -> list[dict] | None:
    """Return tracked-actions table rows as dicts, or None if section not found."""
    in_section = False
    for block in _iter_block_items(document):
        if hasattr(block, "text"):
            if block.text.strip() == _SECTION_HEADING:
                in_section = True
                continue
            if in_section and block.style.name.startswith("Heading"):
                break
        elif in_section and hasattr(block, "rows"):
            return _table_to_dicts(block)
    return None


def _table_to_dicts(table) -> list[dict]:
    headers = [cell.text.strip() for cell in table.rows[0].cells]
    result = []
    for row in table.rows[1:]:
        values = [cell.text.strip() for cell in row.cells]
        if all(v == "" for v in values):
            continue
        result.append(dict(zip(headers, values)))
    return result


def _iter_block_items(document):
    from docx.table import Table
    from docx.text.paragraph import Paragraph
    body = document.element.body
    for child in body:
        if child.tag == qn("w:p"):
            yield Paragraph(child, document)
        elif child.tag == qn("w:tbl"):
            yield Table(child, document)


def find_table_row(document: docx.Document, action_id: int) -> dict | None:
    rows = tracked_actions_table(document)
    if rows is None:
        return None
    for row in rows:
        if str(row.get("ID", "")) == str(action_id):
            return row
    return None
