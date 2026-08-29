"""Offline .docx builder for oracle-parser tests (gts-e5cl).

Builds an in-memory OOXML document from a compact segment spec, so a parser
test can express "this exact paragraph shape" without a live Google Doc.  It
covers the constructs ADR-0027's grammar cares about:

  - soft returns (``\\n`` in a text segment -> ``<w:br/>``) and tabs
    (``\\t`` -> ``<w:tab/>``), the two things ``para.text`` silently drops
  - a PERSON chip (a hyperlink whose URL carries the email)
  - a non-email hyperlink, which is how a flushed ``ACT-N:`` token's chip-badge
    link appears in an export
  - bold runs (ADR-0027 rule 8's field label)
  - list items (``w:numPr``) and table cells, the two containers APT v2 allows

This is a *test fixture* builder, not a serializer: it deliberately shares no
code with ``src/PortableText.js`` or ``scripts/apt_lib.py`` (twin-track
independence — the oracle track must not be able to agree with the
implementation by construction).
"""
import io

import docx
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

_R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_HYPERLINK_RELTYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink"
)


# --- segment constructors -------------------------------------------------

def text(value: str) -> tuple:
    """Plain text.  ``\\n`` becomes a soft return, ``\\t`` becomes a tab."""
    return ("text", value)


def bold(value: str) -> tuple:
    return ("bold", value)


def link(display: str, url: str) -> tuple:
    """A hyperlink whose URL is not an email — e.g. the token's chip badge."""
    return ("link", display, url)


def chip(email: str, display: str) -> tuple:
    """A PERSON chip: a hyperlink whose URL carries the assignee's email."""
    return ("chip", email, display)


def para(*segments, list_item: bool = False) -> dict:
    return {"segments": list(segments), "list_item": list_item}


def table(cells: list[list[list[dict]]]) -> dict:
    """A table: ``cells[r][c]`` is a list of paragraph specs from ``para()``."""
    return {"table": cells}


# --- rendering ------------------------------------------------------------

def _append_text(run_el, value: str) -> None:
    """Append text to a ``w:r``, spelling ``\\n`` as w:br and ``\\t`` as w:tab."""
    buf = ""

    def flush():
        nonlocal buf
        if buf:
            t = OxmlElement("w:t")
            t.text = buf
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            run_el.append(t)
            buf = ""

    for ch in value:
        if ch == "\n":
            flush()
            run_el.append(OxmlElement("w:br"))
        elif ch == "\t":
            flush()
            run_el.append(OxmlElement("w:tab"))
        else:
            buf += ch
    flush()


def _make_run(value: str, *, is_bold: bool = False):
    run = OxmlElement("w:r")
    if is_bold:
        rpr = OxmlElement("w:rPr")
        rpr.append(OxmlElement("w:b"))
        run.append(rpr)
    _append_text(run, value)
    return run


def _add_hyperlink(paragraph, display: str, url: str) -> None:
    r_id = paragraph.part.relate_to(url, _HYPERLINK_RELTYPE, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(f"{{{_R_NS}}}id", r_id)
    hyperlink.append(_make_run(display))
    paragraph._element.append(hyperlink)


def _mark_list_item(paragraph) -> None:
    ppr = paragraph._element.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    ppr.append(num_pr)


def _render_para(paragraph, spec: dict) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)
    if spec.get("list_item"):
        _mark_list_item(paragraph)
    for seg in spec["segments"]:
        kind = seg[0]
        if kind == "text":
            paragraph._element.append(_make_run(seg[1]))
        elif kind == "bold":
            paragraph._element.append(_make_run(seg[1], is_bold=True))
        elif kind == "link":
            _add_hyperlink(paragraph, seg[1], seg[2])
        elif kind == "chip":
            _add_hyperlink(paragraph, seg[2], f"mailto:{seg[1]}")
        else:  # pragma: no cover - programming error in a test spec
            raise ValueError(f"unknown segment kind: {kind!r}")


def build_docx(blocks: list[dict]) -> bytes:
    """Render paragraph/table specs to .docx bytes."""
    doc = Document()
    for block in blocks:
        if "table" in block:
            cells = block["table"]
            rows, cols = len(cells), len(cells[0])
            tbl = doc.add_table(rows=rows, cols=cols)
            for r, row in enumerate(cells):
                for c, cell_paras in enumerate(row):
                    cell = tbl.cell(r, c)
                    first = True
                    for pspec in cell_paras:
                        p = cell.paragraphs[0] if first else cell.add_paragraph()
                        first = False
                        _render_para(p, pspec)
        else:
            _render_para(doc.add_paragraph(), block)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def load(docx_bytes: bytes) -> docx.Document:
    return Document(io.BytesIO(docx_bytes))
